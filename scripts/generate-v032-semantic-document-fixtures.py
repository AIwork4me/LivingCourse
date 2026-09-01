from __future__ import annotations

from pathlib import Path
import argparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
FIXTURE = ROOT / "tests" / "fixtures" / "semantic-manufacturing-course"


def set_font(run, size: float, *, bold: bool = False, color: str = "111111") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size: float, color: str, before: float, after: float) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.25


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Employee Handbook · ")
    set_font(run, 9, color="6B7280")
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.extend([begin, instruction, end])


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, 11)


def make_handbook() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(document.styles["Normal"], 11, "111111", 0, 6)
    configure_style(document.styles["Heading 1"], 16, "2E74B5", 18, 10)
    configure_style(document.styles["Heading 2"], 13, "2E74B5", 14, 7)
    configure_style(document.styles["Heading 3"], 12, "1F4D78", 10, 5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("EMPLOYEE ENTRY HANDBOOK · SYNTHETIC FIXTURE"), 9, bold=True, color="6B7280")
    add_page_number(section.footer.paragraphs[0])

    document.add_paragraph().paragraph_format.space_after = Pt(52)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_font(kicker.add_run("REFERENCE GUIDE"), 10, bold=True, color="2E74B5")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Synthetic Press Entry Handbook"), 28, bold=True, color="203748")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_font(subtitle.add_run("Employee orientation reference · Version 3.0 · Effective 1 September 2026"), 13, color="4B5563")
    purpose = document.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose.paragraph_format.space_before = Pt(72)
    purpose.paragraph_format.space_after = Pt(8)
    set_font(purpose.add_run("For supervised synthetic training only"), 15, bold=True, color="2E74B5")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(note.add_run("This fixture is public-safe and does not describe real equipment."), 11, color="4B5563")
    document.add_page_break()

    document.add_heading("Entry requirements", level=1)
    add_bullet(document, "Confirm that an authorized trainer is present before entering the practice zone.")
    add_bullet(document, "Wear splash goggles and safety shoes before beginning the exercise.")
    add_bullet(document, "Keep a minimum 10 mm clearance from the marked demonstration boundary.")
    add_bullet(document, "Record the simulated pre-start inspection before requesting trainer release.")

    document.add_heading("Quality and escalation", level=1)
    add_bullet(document, "A simulated reading outside ±5% of the training target requires trainer review.")
    add_bullet(document, "If the indicator exceeds 80 °C, stop the exercise and report the observation.")
    add_bullet(document, "When uncertain, stop and ask; do not invent or infer an operating step.")

    document.add_heading("Authority boundary", level=1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    set_font(paragraph.add_run("This handbook is supporting guidance. A controlled SOP overrides it, and any disagreement must be disclosed for author review."), 11, bold=True, color="7A5A00")

    document.add_heading("Workplace note", level=2)
    irrelevant = document.add_paragraph("Lunch breaks may be taken in the west cafeteria between 12:00 and 13:30. This policy is intentionally irrelevant to the course topic.")
    irrelevant.paragraph_format.space_after = Pt(6)

    document.save(FIXTURE / "employee-handbook.docx")


def make_sop() -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("FixtureTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=30, textColor=colors.HexColor("#203748"), alignment=TA_CENTER, spaceAfter=10)
    label = ParagraphStyle("FixtureLabel", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=colors.HexColor("#3D8DFF"), alignment=TA_CENTER, spaceAfter=12)
    h1 = ParagraphStyle("FixtureH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#203748"), spaceBefore=10, spaceAfter=7)
    body = ParagraphStyle("FixtureBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=colors.HexColor("#222222"), spaceAfter=7)
    warning = ParagraphStyle("FixtureWarning", parent=body, fontName="Helvetica-Bold", fontSize=12.5, leading=18, textColor=colors.HexColor("#9B1C1C"), borderColor=colors.HexColor("#E5A8A8"), borderWidth=0.8, borderPadding=9, backColor=colors.HexColor("#FFF5F5"), spaceBefore=5, spaceAfter=10)
    output = FIXTURE / "approved-sop.pdf"
    document = SimpleDocTemplate(str(output), pagesize=letter, rightMargin=0.75 * inch, leftMargin=0.75 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch, title="Approved Synthetic Press SOP")
    story = [
        Paragraph("CONTROLLED INTERNAL · SYNTHETIC PUBLIC-SAFE FIXTURE", label),
        Paragraph("Approved Synthetic Press SOP", title),
        Paragraph("Owner: Fixture Safety Office · Version 4.2 · Effective 1 September 2026", body),
        Spacer(1, 0.12 * inch),
    ]
    metadata = Table([
        ["Authority", "controlled_internal"],
        ["Approval", "Fixture Safety Owner — approved"],
        ["Scope", "Supervised software-validation exercise only"],
    ], colWidths=[1.35 * inch, 5.65 * inch])
    metadata.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#203748")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8BCC4")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.extend([
        metadata,
        Spacer(1, 0.18 * inch),
        Paragraph("1. Entry and protective equipment", h1),
        Paragraph("An authorized trainer must be present before the practice zone is opened. Wear splash goggles and safety shoes before starting.", body),
        Paragraph("2. Controlled synthetic setting", h1),
        Paragraph("Synthetic training pressure setting = 0.55 MPa", warning),
        Paragraph("Use this fictional value only inside the synthetic fixture. It intentionally conflicts with an archived training deck so authority resolution can be tested.", body),
        Paragraph("3. Guard and warning controls", h1),
        Paragraph("禁止在设备运行时打开防护门。 Do not open the guard door while the synthetic machine is running.", warning),
        Paragraph("If the simulated warning beacon illuminates, stop the exercise and report the event to the trainer.", body),
        Paragraph("4. Pre-start sequence", h1),
        Paragraph("Confirm the guard is closed; verify the training label; record the simulated pre-start inspection; then wait for trainer release.", body),
        PageBreak(),
        Paragraph("Evidence continuation", h1),
        Paragraph("The second page gives the fixture a stable page-level evidence location. It introduces no replacement pressure setting.", body),
        Paragraph("Device-grounding limitation", h1),
        Paragraph("This controlled text does not confirm a real device image, component anchor, control point, or operation region. Production use remains blocked until responsible human review supplies that grounding.", body),
    ])
    document.build(story)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("artifact", choices=("docx", "pdf"))
    args = parser.parse_args()
    FIXTURE.mkdir(parents=True, exist_ok=True)
    if args.artifact == "docx":
        make_handbook()
    else:
        make_sop()
