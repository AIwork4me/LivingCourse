from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
FIXTURE = ROOT / "tests" / "fixtures" / "raw-manufacturing-course"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def make_handbook() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(30)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(24, 48, 65)

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("SYNTHETIC PUBLIC-SAFE TEST MATERIAL")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(39, 126, 170)
    title = document.add_paragraph("Employee Entry Handbook", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Synthetic Training Machine · New-hire orientation reference")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    document.add_paragraph()
    summary = document.add_table(rows=3, cols=2)
    summary.style = "Light Shading Accent 1"
    for row, values in zip(summary.rows, [("Owner", "Synthetic HR Training"), ("Version", "Fixture 1.0"), ("Effective", "1 September 2026")]):
        row.cells[0].text, row.cells[1].text = values
    document.add_page_break()

    document.add_heading("Before entering the practice area", level=1)
    document.add_paragraph("This handbook supports author-review testing only. It is not a real operating instruction and contains no real equipment settings.")
    for item in [
        "Confirm the trainer is present before starting a simulated activity.",
        "Wear the PPE named in the current controlled site procedure.",
        "Stop the exercise when the simulated warning beacon appears.",
        "Report uncertainty instead of guessing an operating step.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Authority and escalation", level=1)
    document.add_paragraph("This handbook is a reference source. When it differs from a controlled internal SOP, the discrepancy must be shown to an author. LivingCourse must not silently select a ground truth.")
    document.add_heading("Device evidence limitation", level=1)
    document.add_paragraph("The phrase Synthetic Training Machine describes a fictional training prop. The fixture photo is illustrative and does not confirm a real device, component, control point, or operation region.")
    add_page_number(section.footer.paragraphs[0])
    document.save(FIXTURE / "employee-handbook.docx")


def make_sop() -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("FixtureTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, textColor=colors.HexColor("#183041"), alignment=TA_CENTER, spaceAfter=10)
    label_style = ParagraphStyle("FixtureLabel", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=colors.HexColor("#277EAA"), alignment=TA_CENTER, spaceAfter=12)
    h1 = ParagraphStyle("FixtureH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#183041"), spaceBefore=10, spaceAfter=7)
    body = ParagraphStyle("FixtureBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=colors.HexColor("#222222"), spaceAfter=7)
    output = FIXTURE / "sop.pdf"
    doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=17 * mm, bottomMargin=17 * mm, title="Synthetic Training Machine Controlled SOP")
    story = [
        Paragraph("SYNTHETIC PUBLIC-SAFE TEST MATERIAL", label_style),
        Paragraph("Controlled Practice SOP", title_style),
        Paragraph("Synthetic Training Machine · Document owner: Fixture Safety Owner", body),
        Spacer(1, 4 * mm),
    ]
    table = Table([
        ["Source class", "controlled_internal"],
        ["Version", "2.0"],
        ["Effective date", "1 September 2026"],
        ["Scope", "Simulated author-review evaluation only"],
    ], colWidths=[42 * mm, 118 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF5FB")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#183041")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8BCC4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([
        table,
        Spacer(1, 6 * mm),
        Paragraph("1. Controlled synthetic setting", h1),
        Paragraph("Synthetic training pressure setting = B", ParagraphStyle("Setting", parent=body, fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#0B6E99"), borderColor=colors.HexColor("#B8BCC4"), borderWidth=0.8, borderPadding=9, backColor=colors.HexColor("#F4FAFD"))),
        Paragraph("The letter B is a fictional comparison token. It is intentionally incompatible with the old deck's letter A so the conflict resolver can be tested without publishing a hazardous real-world parameter.", body),
        Paragraph("2. Simulated exercise sequence", h1),
        Paragraph("1. Confirm the trainer has opened the practice session.<br/>2. Select the fictional setting B on the training card.<br/>3. Stop when the simulated warning beacon appears.<br/>4. Record the observation and wait for trainer release.", body),
        Paragraph("3. Production limitation", h1),
        Paragraph("This SOP is synthetic evidence for software validation. It does not identify a real device or a confirmed operation region. Production training must remain blocked until a real approved source and real-device anchor are reviewed by a responsible human.", body),
        PageBreak(),
        Paragraph("Evidence-location test page", h1),
        Paragraph("This second page exists to validate page grouping and page-level EvidenceRef locations. It repeats no setting value.", body),
    ])
    doc.build(story)


def make_photo() -> None:
    width, height = 1280, 800
    image = Image.new("RGB", (width, height), "#EDF5F8")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((120, 120, 1160, 680), radius=36, fill="#D5E5EA", outline="#648997", width=8)
    draw.rounded_rectangle((230, 210, 760, 600), radius=24, fill="#F8FBFC", outline="#648997", width=6)
    draw.rectangle((810, 245, 1030, 530), fill="#B9D0D8", outline="#486A76", width=6)
    draw.ellipse((865, 310, 975, 420), fill="#6DCBF4", outline="#277EAA", width=8)
    draw.line((180, 635, 1100, 635), fill="#486A76", width=10)
    draw.polygon([(760, 380), (830, 340), (830, 420)], fill="#FFB65C", outline="#8B5A1D")
    image.save(FIXTURE / "equipment-photo.jpg", quality=92, optimize=True)


if __name__ == "__main__":
    FIXTURE.mkdir(parents=True, exist_ok=True)
    make_handbook()
    make_sop()
    make_photo()
